import numpy as np
import warnings
import pandas as pd
from openpyxl import load_workbook, Workbook
import os
import docx
from tqdm import tqdm


def ohms_law(voltage=None, resistance=None, current=None):
    if voltage is not None and resistance is not None:
        return np.array(voltage, dtype=np.float64) / np.array(resistance, dtype=np.float64)
    elif voltage is not None and current is not None:
        return np.array(voltage, dtype=np.float64) / np.array(current, dtype=np.float64)
    elif current is not None and resistance is not None:
        return np.array(current, dtype=np.float64) * np.array(resistance, dtype=np.float64)
    else:
        warnings.warn("Id10 Error: not given enough knows.")
        return np.array([])


def resistors_parallel(resistors):
    resistors = np.array(resistors, dtype=np.float64)
    if np.any(resistors == 0):
        return 0
    return 1 / (1 / resistors).sum()


def resistors_series(resistors):
    return np.array(resistors, dtype=np.float64).sum()


def capacitor_parallel(caps):
    return np.array(caps, dtype=np.float64).sum()


def capacitor_series(cap):
    cap = np.array(cap, dtype=np.float64)
    if np.any(cap == 0):
        return 0
    return 1 / (1 / cap).sum()


def percent_difference(measured, calculated):
    measured = np.array(measured, dtype=np.float64)
    calculated = np.array(calculated, dtype=np.float64)
    return np.abs((calculated - measured) / calculated) * 100


def current_division(resistors, current=None, voltage=None):
    resistors = np.array(resistors, dtype=np.float64)
    df = pd.DataFrame()
    df["resistors"] = resistors
    if voltage is not None:
        df["current"] = voltage / resistors
    elif current is not None:
        df["current"] = ohms_law(current=current,  resistance=resistors_parallel(resistors)) / resistors
    else:
        warnings.warn("Id10 Error: missing both current and voltage. One must be set for calculation to work")
        return None
    return df


def voltage_division(resistors, voltage=None, current=None):
    resistors = np.array(resistors)
    df = pd.DataFrame()
    df["resistors"] = resistors
    if current is not None:
        df["voltage"] = ohms_law(resistance=resistors, current=current)
    elif voltage is not None:
        df["voltage"] = ohms_law(resistance=resistors, current=ohms_law(resistance=resistors.sum(), voltage=voltage))
    else:
        warnings.warn("Id10 Error: missing both current and voltage. One must be set for calculation to work")
        return None
    return df


def resistor_tolerance(resistors, tolerances):
    resistors = np.array(resistors, dtype=np.float64)
    tolerances = np.array(tolerances, dtype=np.float64)
    diff = resistors * tolerances
    df = pd.DataFrame()
    df["resistors"] = resistors
    df["tolerances"] = tolerances
    df["max_in_spec"] = resistors + diff
    df["min_in_spec"] = resistors - diff
    return df


def resistor_in_spec(resistor, measure, tolerance):
    res = np.array(resistor, dtype=np.float64)
    mea = np.array(measure, dtype=np.float64)
    tol = np.array(tolerance, dtype=np.float64)
    df = pd.DataFrame([], dtype=str)
    df['Normal'] = res
    df['Measured'] = mea
    df['In Spec'] = np.where(abs(res-mea) < res*tol, "Yes", "No")
    df['% Diff'] = np.array(np.abs((res - mea) / res)*100, dtype=np.float64)
    return df


def save_data_to_exec(dataframes, file_name):
    if not os.path.exists(file_name):
        wb = Workbook(write_only=True)
        wb.save(file_name)
        wb.close()
        book = load_workbook(file_name, read_only=False)
        book.remove(book["Sheet"])
    else:
        book = load_workbook(file_name, read_only=False)
    writer = pd.ExcelWriter(file_name, engine='openpyxl')
    writer.book = book
    for df in tqdm(dataframes, unit="DataFrame", ascii=True, desc="Saving DataFrames to excel!"):
        df[1].to_excel(writer, sheet_name=df[0])
    writer.save()
    writer.close()
    print("\n"*3)


def save_data_to_word_table(dataframes, file_name):
    if not os.path.exists(file_name):
        doc = docx.Document()
    else:
        doc = docx.Document(file_name)
    for df in tqdm(dataframes, unit="DataFrame", ascii=True, desc="Saving DataFrames to tables in word!"):
        doc.add_paragraph(df[0])
        frame = df[1]
        table = doc.add_table(rows=1, cols=frame.shape[1], style='TableGrid')
        header_cells = table.rows[0].cells
        for cell in tqdm(range(len(header_cells)), unit="Cell", ascii=True, desc="Adding Headers to table: {}".format(
                df[0])):
            header_cells[cell].text = frame.columns[cell]
        for row in tqdm(range(frame.shape[0]), unit="Row", ascii=True, desc="Adding Rows to table: {}".format(df[0])):
            cells = table.add_row().cells
            for cell in tqdm(range(len(cells)), unit="Cell", ascii=True, desc="Adding date to row"):
                cells[cell].text = str(frame.values[row, cell])
        doc.add_paragraph("\n\n\n")
    doc.save(file_name)
    print("\n"*4)


def load_data_from_exec(file_name):
    data = []
    if os.path.exists(file_name):
        excel = pd.ExcelFile(file_name)
        for sheet in excel.sheet_names:
            df = excel.parse(sheet)
            if not df.empty:
                data.append([sheet, df])
    return data


def mesh_analysis(resistances, voltages):
    resistances = np.matrix(resistances)
    voltages = np.matrix(voltages)
    if resistances.shape[0] != voltages.shape[0]:
        voltages = np.reshape(voltages, (resistances.shape[0], 1))
    return np.linalg.inv(resistances) * voltages


def capacitor_impedance(capacitor, frequency):
    np.seterr(divide='ignore', invalid='ignore')
    capacitor = np.array(capacitor, dtype=np.float64)
    frequency = np.array(frequency, dtype=np.float64)
    return np.divide(-1j, (capacitor * frequency))


def inductor_impedance(inductor, frequency):
    inductor = np.array(inductor, dtype=np.float64)
    frequency = np.array(frequency, dtype=np.float64)
    return inductor * frequency * 1j

